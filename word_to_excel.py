import os
import shutil
import json
import pandas as pd
from docx import Document
import re
import logging
import time
import threading
import sys
import msvcrt

def load_config(config_path):
    with open(config_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def setup_logger(log_path):
    logging.basicConfig(
        filename=log_path,
        filemode='a',
        format='%(asctime)s %(levelname)s: %(message)s',
        level=logging.INFO,
        encoding='utf-8'
    )

def load_symbol_config(config_path):
    # 新增：加载符号映射和tick_symbols配置
    with open(config_path, 'r', encoding='utf-8') as f:
        config = json.load(f)
    symbol_maps = config.get('symbol_maps', {})
    tick_symbols = config.get('tick_symbols', ['✓', '✔', '☑', '☒', '√'])
    empty_box = config.get('empty_box', '☐')
    return symbol_maps, tick_symbols, empty_box

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    return '\n'.join([p.text for p in doc.paragraphs])
#解决word中wingdings123对号无法识别
def extract_text_with_symbols(docx_path, symbol_maps):
    doc = Document(docx_path)
    lines = []
    # 读取配置的映射表
    wingdings_map = symbol_maps.get('wingdings', {})
    wingdings2_map = symbol_maps.get('wingdings2', {})
    wingdings3_map = symbol_maps.get('wingdings3', {})
    for para in doc.paragraphs:
        line = ''
        for run in para.runs:
            text = run.text
            font = run.font.name or ''
            orig_text = text
            # Wingdings 1
            if 'Wingdings' == font:
                for k, v in wingdings_map.items():
                    if k in text:
                        text = text.replace(k, v)
            # Wingdings 2
            elif 'Wingdings 2' == font:
                for k, v in wingdings2_map.items():
                    if k in text:
                        text = text.replace(k, v)
            # Wingdings 3
            elif 'Wingdings 3' == font:
                for k, v in wingdings3_map.items():
                    if k in text:
                        text = text.replace(k, v)
            line += text
        lines.append(line)
    result = '\n'.join(lines)
    return result

def format_date_to_str(text):
    # 匹配如2016.01、2016.10、2016-01、2016/01等，统一转为'2016.01'
    def repl(m):
        y, mth = m.group(1), m.group(2)
        return f"'{y}.{mth.zfill(2)}"  # 前导单引号防止Excel自动转格式
    return re.sub(r'(\d{4})[./-](\d{1,2})', repl, text)

def monitor_keyboard(pause_flag, exit_flag):
    while True:
        key = msvcrt.getch()
        if key == b' ':
            pause_flag[0] = not pause_flag[0]
            if pause_flag[0]:
                logging.info("检测到空格键，处理暂停。再次按空格键恢复。")
                print("[提示] 已暂停，按空格键恢复。")
            else:
                logging.info("检测到空格键，恢复处理。")
                print("[提示] 已恢复，继续处理。")
        elif key == b'\x1b':  # ESC
            exit_flag[0] = True
            logging.info("检测到ESC键，程序即将退出。")
            print("[提示] 检测到ESC键，程序即将退出。")
            break

def main():
    config = load_config('config.json')
    src = config['source_folder']
    backup = config['backup_folder']
    excel_path = config['excel_path']
    keywords = config['keywords']
    log_path = config.get('log_path', './process.log')
    setup_logger(log_path)

    pause_flag = [False]
    exit_flag = [False]
    kb_thread = threading.Thread(target=monitor_keyboard, args=(pause_flag, exit_flag), daemon=True)
    kb_thread.start()

    try:
        logging.info(f"程序启动，配置内容：{config}")
        print("配置内容：", config)
        print("源目录：", src)
        print("目录下文件：", os.listdir(src))
    except Exception as e:
        logging.error(f"初始化异常: {e}")
        return

    # 输出文件保存到“结果”文件夹
    result_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '结果')
    os.makedirs(result_dir, exist_ok=True)
    out_filename = os.path.basename(excel_path)
    excel_path = os.path.join(result_dir, out_filename)

    # 读取或新建Excel/CSV
    if excel_path.lower().endswith('.csv'):
        if os.path.exists(excel_path):
            try:
                df = pd.read_csv(excel_path, dtype=str, encoding='utf-8-sig')
            except UnicodeDecodeError:
                df = pd.read_csv(excel_path, dtype=str, encoding='gbk')
        else:
            df = pd.DataFrame({k: [] for k in keywords}, dtype=str)
    else:
        if os.path.exists(excel_path):
            df = pd.read_excel(excel_path, dtype=str, engine='openpyxl')
        else:
            df = pd.DataFrame({k: [] for k in keywords}, dtype=str)

    total_written = 0
    symbol_maps, tick_symbols, empty_box = load_symbol_config('config.json')
    while True:
        if exit_flag[0]:
            msg = "[提示] 程序已退出。"
            print(msg)
            logging.info(msg)
            break
        if pause_flag[0]:
            msg = "[提示] 暂停中，按空格键恢复。"
            print(msg)
            logging.info(msg)
            while pause_flag[0] and not exit_flag[0]:
                time.sleep(2)
            if exit_flag[0]:
                msg = "[提示] 程序已退出。"
                print(msg)
                logging.info(msg)
                break
        file_list = [f for f in os.listdir(src) if f.lower().endswith('.docx') and not f.startswith('~$')]
        if not file_list:
            time.sleep(2)
            continue
        for fname in file_list:
            fpath = os.path.join(src, fname)
            # 检查文件是否正在被复制（大小稳定才处理）
            try:
                size1 = os.path.getsize(fpath)
                time.sleep(0.5)
                size2 = os.path.getsize(fpath)
                if size1 != size2:
                    continue  # 文件还在复制中，跳过
            except Exception as e:
                logging.warning(f"文件 {fname} 检查大小异常: {e}")
                continue
            try:
                text = extract_text_with_symbols(fpath, symbol_maps)
                text = format_date_to_str(str(text))
                lines = text.split('\n')
                row = {k: '' for k in keywords}
                for idx, line in enumerate(lines):
                    kw_matches = []
                    for kw in keywords:
                        kw_pattern = r'{}[ ]*[：:][ ]*'.format(re.escape(kw))
                        for m in re.finditer(kw_pattern, line):
                            kw_matches.append({'kw': kw, 'start': m.start(), 'end': m.end()})
                    kw_matches.sort(key=lambda x: x['start'])
                    for i, match in enumerate(kw_matches):
                        kw = match['kw']
                        start = match['end']
                        end = kw_matches[i+1]['start'] if i+1 < len(kw_matches) else len(line)
                        val = line[start:end].strip()
                        if any(sym in val for sym in tick_symbols) or empty_box in val:
                            stack = []
                            i2 = 0
                            while i2 < len(val):
                                if val[i2] == empty_box:
                                    j = i2 + 1
                                    while j < len(val) and val[j] not in tick_symbols + [empty_box]:
                                        j += 1
                                    content = val[i2+1:j].strip()
                                    if content:
                                        stack.append({'type': 'empty', 'content': content, 'pos': i2})
                                    i2 = j
                                elif val[i2] in tick_symbols:
                                    j = i2 + 1
                                    while j < len(val) and val[j] not in tick_symbols + [empty_box]:
                                        j += 1
                                    content = val[i2+1:j].strip()
                                    if content:
                                        stack.append({'type': 'tick', 'content': content, 'pos': i2})
                                    i2 = j
                                else:
                                    i2 += 1
                            tick_indices = [ii for ii, item in enumerate(stack) if item['type'] == 'tick']
                            if tick_indices:
                                first_tick = tick_indices[0]
                                stack = stack[first_tick:]
                            val = ' '.join([item['content'] for item in stack if item['type'] == 'tick']).strip()
                            if val == '' or val.startswith(' '):
                                val = ''
                        row[kw] = str(val)
                for kw in keywords:
                    if kw not in row or row[kw] == '':
                        row[kw] = ''
                df = pd.concat([df, pd.DataFrame([row])], ignore_index=True)
                total_written += 1
                os.makedirs(backup, exist_ok=True)
                shutil.move(fpath, os.path.join(backup, fname))
                logging.info(f"处理文件: {fname}，插入内容: {row}")
            except Exception as e:
                logging.error(f"处理文件 {fname} 异常: {e}")
        # 写入Excel/CSV
        try:
            df = df.astype(str)
            if excel_path.lower().endswith('.csv'):
                df.to_csv(excel_path, index=False, encoding='utf-8-sig')
            else:
                df.to_excel(excel_path, index=False, engine='openpyxl')
            logging.info(f"共写入 {total_written} 条数据到Excel/CSV。实际输出路径：{excel_path}")
        except Exception as e:
            logging.error(f"写入Excel/CSV异常: {e}")
        time.sleep(2)

if __name__ == '__main__':
    try:
        main()
    except Exception as e:
        logging.error(f"主程序异常: {e}")
